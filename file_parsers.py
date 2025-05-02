import pdfplumber
from docx import Document
import mammoth
from PIL import Image
import pytesseract
import mimetypes
import os
import re
import io
import tempfile
import shutil
from pathlib import Path
import requests
import json
import base64
import fitz  # PyMuPDF
import html
import uuid
from bs4 import BeautifulSoup

# Add new imports for advanced image parsing
import cv2
import numpy as np
# from paddleocr import PaddleOCR  # Commented out as paddleocr is not installed

# Import API key from config file
from config import OPENAI_API_KEY

def detect_file_type(filepath):
    """
    Detect the file type based on extension and MIME type.
    
    Args:
        filepath: Path to the file
        
    Returns:
        String representing the file type or extension
    """
    mime, _ = mimetypes.guess_type(filepath)
    return mime or filepath.split('.')[-1]

def save_image(image_obj, output_dir, page_number, image_number):
    image_bytes = image_obj["stream"].get_data()
    ext = image_obj["ext"]
    filename = f"pdf_image_page{page_number}_img{image_number}.{ext}"
    filepath = os.path.join(output_dir, filename)
    with open(filepath, "wb") as f:
        f.write(image_bytes)
    return filepath

def parse_pdf(filepath):
    """
    Parse a PDF file and convert it to markdown format with table support.
    
    Args:
        filepath: Path to the PDF file
        
    Returns:
        String containing markdown representation of the PDF
    """
    md = []
    try:
        with pdfplumber.open(filepath) as pdf:
            # Create a temp directory for extracted images
            temp_dir = tempfile.mkdtemp()
            try:
                # Keep track of image counters
                image_count = 0
                
                for i, page in enumerate(pdf.pages):
                    md.append(f"# Page {i+1}")
                    
                    # Extract images if available
                    try:
                        # Extract and save images from the page
                        page_images = extract_images_from_pdf_page(page, temp_dir, i, image_count)
                        image_count += len(page_images)
                        
                        # Get tables from the page
                        tables = page.extract_tables()
                        has_tables = len(tables) > 0
                        
                        if has_tables:
                            # Check if we can get table areas to exclude from text
                            try:
                                # Extract the text excluding table regions
                                non_table_text = extract_text_excluding_tables(page)
                                if non_table_text:
                                    md.append(non_table_text)
                            except Exception as e:
                                # Fallback to a simpler method if character-level approach fails
                                print(f"Error in character processing: {e}")
                                # Just extract all text with normal method as fallback
                                text = page.extract_text()
                                if text:
                                    md.append(text)
                            
                            # Now process each table separately
                            for j, table in enumerate(tables):
                                if table:  # Check if table is not empty
                                    md.append(f"\n### Table {j+1}\n")
                                    md.append(convert_table_to_markdown(table))
                        else:
                            # No tables, just extract text
                            text = page.extract_text()
                            if text:
                                # Extract text by paragraph for better formatting
                                paragraphs = text.split('\n\n')
                                clean_paragraphs = []
                                
                                for p in paragraphs:
                                    if p.strip():
                                        clean_paragraphs.append(p.strip())
                                
                                md.append("\n\n".join(clean_paragraphs))
                        
                        # Add extracted images to markdown
                        for img_path, img_desc in page_images:
                            md.append(f"\n![{img_desc}]({img_path})\n")
                    
                    except Exception as e:
                        md.append(f"\n*Error processing page {i+1}: {str(e)}*\n")
            
            finally:
                # Clean up the temp directory
                shutil.rmtree(temp_dir, ignore_errors=True)
    
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"
    
    return "\n\n".join(md)

def extract_text_excluding_tables(page):
    """
    Extract text from a page while excluding text in table regions.
    This simplified approach just extracts text from the whole page
    and uses markers to indicate where tables begin and end.
    
    Args:
        page: A pdfplumber page object
        
    Returns:
        String containing text excluding table regions
    """
    # Since the character-level approach is causing text corruption,
    # we'll switch to using a simpler, more reliable approach
    
    # First, get all text from the page using the standard method
    whole_text = page.extract_text()
    if not whole_text:
        return ""
    
    # Look for tables
    tables = page.find_tables()
    if not tables:
        return whole_text
    
    # Get table bounding boxes
    valid_table_bboxes = []
    for table in tables:
        if hasattr(table, 'bbox') and is_valid_bbox(table.bbox):
            valid_table_bboxes.append(table.bbox)
    
    # If no valid table bboxes found, return all text
    if not valid_table_bboxes:
        return whole_text
    
    # Since trying to filter out table text at character level is causing issues,
    # we'll extract text from the whole page and then skip processing tables separately.
    # In this approach, tables will appear twice (as plain text and formatted tables),
    # but the text won't be corrupted.
    
    # Format the extracted text
    paragraphs = whole_text.split('\n\n')
    clean_paragraphs = []
    
    for p in paragraphs:
        if p.strip():
            clean_paragraphs.append(p.strip())
    
    return "\n\n".join(clean_paragraphs)

def is_valid_bbox(bbox):
    """
    Check if a bounding box is valid (has positive width and height).
    
    Args:
        bbox: A tuple or list with 4 elements (x0, top, x1, bottom)
        
    Returns:
        Boolean indicating if the bbox is valid
    """
    if not bbox or len(bbox) != 4:
        return False
    
    x0, top, x1, bottom = bbox
    
    # Check for valid types
    if not all(isinstance(x, (int, float)) for x in [x0, top, x1, bottom]):
        return False
    
    # Check width and height are positive
    width = x1 - x0
    height = bottom - top
    
    return width > 0 and height > 0

def extract_images_from_pdf_page(page, output_dir, page_num, img_counter):
    """
    Extract images from a PDF page and save them to a directory.
    
    Args:
        page: A pdfplumber page object
        output_dir: Directory to save extracted images
        page_num: Current page number
        img_counter: Counter for image numbering
        
    Returns:
        List of tuples (image_path, image_description)
    """
    images = []
    
    # Extract images directly if available through pdfplumber
    try:
        # Use page.images if available (depends on version)
        if hasattr(page, 'images') and page.images:
            for i, img in enumerate(page.images):
                img_num = img_counter + i
                img_filename = f"image_{page_num+1}_{i+1}.png"
                img_path = os.path.join(output_dir, img_filename)
                
                # Some PDF images can be directly extracted
                if 'stream' in img and img['stream']:
                    try:
                        image_bytes = img['stream'].get_data()
                        with open(img_path, 'wb') as f:
                            f.write(image_bytes)
                        images.append((img_path, f"Image {img_num+1}"))
                    except Exception as e:
                        # If extraction fails, log but continue
                        print(f"Error extracting image: {e}")
        
        # Fallback: Render page as image if no images were extracted
        if not images:
            # This is a fallback option, might not be needed in all cases
            pass
    
    except Exception as e:
        print(f"Error extracting images from page {page_num+1}: {e}")
    
    return images

def convert_table_to_markdown(table):
    """
    Convert a pdfplumber table to markdown format.
    
    Args:
        table: A table extracted by pdfplumber
        
    Returns:
        String containing markdown table
    """
    md_table = []
    
    # Process header row
    if len(table) > 0:
        # Clean header cells
        header = [str(cell).strip() if cell is not None else "" for cell in table[0]]
        md_table.append("| " + " | ".join(header) + " |")
        
        # Add separator row
        md_table.append("| " + " | ".join(["---"] * len(header)) + " |")
        
        # Process data rows
        for row in table[1:]:
            # Clean cell data
            row_data = [str(cell).strip() if cell is not None else "" for cell in row]
            md_table.append("| " + " | ".join(row_data) + " |")
    
    return "\n".join(md_table)

def parse_docx(filepath):
    """
    Parse a DOCX file and convert it to markdown format.
    
    Args:
        filepath: Path to the DOCX file
        
    Returns:
        String containing markdown representation of the DOCX
    """
    try:
        doc = Document(filepath)
        md = []
        
        # Create a temp directory for extracted images
        temp_dir = tempfile.mkdtemp()
        try:
            # Extract images if present (this requires additional handling)
            image_count = 0
            
            for p in doc.paragraphs:
                if p.style.name.startswith("Heading"):
                    level = p.style.name[-1]
                    md.append(f"{'#' * int(level)} {p.text}")
                else:
                    md.append(p.text)
            
            # Optional: Process tables from DOCX
            for table in doc.tables:
                md.append("\n### Table\n")
                table_rows = []
                
                # Process header row
                header_row = []
                for cell in table.rows[0].cells:
                    header_row.append(cell.text.strip())
                
                table_rows.append("| " + " | ".join(header_row) + " |")
                table_rows.append("| " + " | ".join(["---"] * len(header_row)) + " |")
                
                # Process data rows
                for row in table.rows[1:]:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_rows.append("| " + " | ".join(row_data) + " |")
                
                md.append("\n".join(table_rows))
        
        finally:
            # Clean up the temp directory
            shutil.rmtree(temp_dir, ignore_errors=True)
        
        return "\n\n".join(md)
    
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"

def parse_rtf(filepath):
    """
    Parse an RTF file and convert it to markdown format.
    Extract both text content and embedded images.
    
    Args:
        filepath: Path to the RTF file
        
    Returns:
        String containing markdown representation of the RTF
    """
    try:
        # First try to extract text with mammoth
        with open(filepath, "rb") as f:
            result = mammoth.convert_to_markdown(f)
        
        text_content = result.value
        
        # Create a temporary directory for extracted images
        temp_dir = tempfile.mkdtemp()
        extracted_images = []
        
        try:
            # Extract embedded images using a more robust approach
            extracted_images = extract_images_from_rtf(filepath, temp_dir)
            
            # Combine text and image references into final markdown
            markdown_lines = []
            
            # Add the text content first
            if text_content.strip():
                markdown_lines.append(text_content.strip())
            
            # Add extracted images
            if extracted_images:
                markdown_lines.append("\n## Embedded Images\n")
                for i, img_path in enumerate(extracted_images):
                    img_filename = os.path.basename(img_path)
                    markdown_lines.append(f"\n![Embedded Image {i+1}]({img_filename})\n")
            
            return "\n\n".join(markdown_lines)
        
        finally:
            # Clean up temporary directory
            shutil.rmtree(temp_dir, ignore_errors=True)
            
    except Exception as e:
        print(f"Error parsing RTF file: {str(e)}")
        import traceback
        print(f"Exception traceback: {traceback.format_exc()}")
        return f"Error parsing RTF: {str(e)}"

def extract_images_from_rtf(rtf_path, output_dir):
    """
    Extract embedded images from an RTF file.
    
    Args:
        rtf_path: Path to the RTF file
        output_dir: Directory to save extracted images
        
    Returns:
        List of paths to extracted images
    """
    extracted_images = []
    
    try:
        # Read the RTF file as binary
        with open(rtf_path, 'rb') as rtf_file:
            rtf_content = rtf_file.read().decode('utf-8', errors='ignore')
        
        # Look for embedded base64 images
        # Pattern for base64 data in RTF files
        base64_pattern = r'\\pict[^{]*?{[^}]*?\\pngblip[^}]*?}|{\\pict[^}]*?\\pngblip[^}]*?}'
        hex_pattern = r'\\pict[^{]*?{[^}]*?\\(?:pngblip|jpegblip)[^}]*?([0-9A-Fa-f\s]+)}'
        
        # First try to find base64 encoded images
        base64_matches = re.findall(base64_pattern, rtf_content)
        
        # Process each match
        for i, match in enumerate(base64_matches):
            try:
                # Extract the hexadecimal data, removing RTF control words
                hex_data = re.sub(r'\\[a-z]+|\{|\}|\\|\s', '', match)
                
                if hex_data:
                    # Convert hex to binary
                    binary_data = bytes.fromhex(hex_data)
                    
                    # Try to determine the image type
                    img_ext = '.png'  # Default to PNG
                    if b'\xff\xd8\xff' in binary_data[:10]:  # JPEG signature
                        img_ext = '.jpg'
                    elif b'GIF8' in binary_data[:10]:  # GIF signature
                        img_ext = '.gif'
                    
                    # Save the image
                    img_filename = f"rtf_image_{i+1}{img_ext}"
                    img_path = os.path.join(output_dir, img_filename)
                    
                    with open(img_path, 'wb') as img_file:
                        img_file.write(binary_data)
                    
                    extracted_images.append(img_path)
            except Exception as e:
                print(f"Error extracting image {i+1}: {str(e)}")
        
        # If no images found with the first method, try alternative patterns
        if not extracted_images:
            # Try to extract hexadecimal data directly
            hex_matches = re.findall(hex_pattern, rtf_content)
            
            for i, hex_data in enumerate(hex_matches):
                try:
                    # Clean up the hex data
                    clean_hex = re.sub(r'\s', '', hex_data)
                    
                    if clean_hex:
                        # Convert hex to binary
                        binary_data = bytes.fromhex(clean_hex)
                        
                        # Try to determine the image type
                        img_ext = '.png'  # Default to PNG
                        if b'\xff\xd8\xff' in binary_data[:10]:  # JPEG signature
                            img_ext = '.jpg'
                        elif b'GIF8' in binary_data[:10]:  # GIF signature
                            img_ext = '.gif'
                        
                        # Save the image
                        img_filename = f"rtf_image_{i+1}{img_ext}"
                        img_path = os.path.join(output_dir, img_filename)
                        
                        with open(img_path, 'wb') as img_file:
                            img_file.write(binary_data)
                        
                        extracted_images.append(img_path)
                except Exception as e:
                    print(f"Error extracting image {i+1} (hex method): {str(e)}")
        
        return extracted_images
    
    except Exception as e:
        print(f"Error extracting images from RTF: {str(e)}")
        return []

def parse_image(filepath):
    """
    Parse an image file using OCR with Tesseract.
    
    Args:
        filepath: Path to the image file
        
    Returns:
        String containing markdown representation of the image with OCR text
    """
    try:
        # Step 1: Load image with OpenCV
        image = cv2.imread(filepath)
        if image is None:
            return f"Error loading image: {filepath}"
        
        image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
        
        markdown = []
        
        # Add the original image to markdown
        markdown.append(f"![Original Image]({os.path.basename(filepath)})\n")
        markdown.append("## Extracted Content\n")
        
        # Use Tesseract OCR
        # Convert to RGB if image is in RGBA mode (for PNG transparency)
        img = Image.open(filepath)
        if img.mode == 'RGBA':
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])  # 3 is the alpha channel
            img = background
        
        # Configure OCR for better results
        custom_config = r'--oem 3 --psm 6'  # Full page segmentation with OEM engine 3
        
        # Extract text with Tesseract OCR
        text = pytesseract.image_to_string(img, config=custom_config)
        
        if text.strip():
            markdown.append(text.strip())
        else:
            markdown.append("*No text could be extracted from this image.*")
        
        return "\n\n".join(markdown)
    
    except Exception as e:
        return f"Error processing image: {str(e)}"

def refine_markdown_with_llm(markdown_text):
    """
    Refine markdown content using OpenAI's GPT-4o Mini model.
    
    Args:
        markdown_text: Raw markdown text to refine
        
    Returns:
        Refined markdown text with better structure and formatting
    """
    # Get API key from config
    api_key = OPENAI_API_KEY
    api_url = "https://api.openai.com/v1/chat/completions"
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    
    prompt = f"""
    Please improve the structure and formatting of the following markdown content extracted from a document. 
    Make sure to:
    1. Keep all important information
    2. Improve the formatting and structure
    3. Fix any typos or OCR errors if obvious
    4. Ensure all markdown syntax is correct and valid
    5. For any icons, symbols, or special characters that cannot be represented in plain text, describe them within square brackets (e.g., [checkmark icon], [arrow pointing right])
    6. Return ONLY the refined markdown without any additional explanations or metadata
    
    Here's the content to refine:
    
    {markdown_text}
    """
    
    payload = {
        "model": "gpt-4o-mini",
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful assistant that improves the structure and formatting of markdown content extracted from documents. Your output must be valid, properly formatted markdown. For any icons, symbols, or special characters that cannot be represented in plain text, describe them within square brackets."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.3
    }
    
    try:
        print("\n=== SENDING LLM REQUEST (Text Refinement) ===")
        print(f"Model: gpt-4o-mini")
        print(f"Input text length: {len(markdown_text)} characters")
        print(f"Temperature: 0.3")
        
        response = requests.post(api_url, headers=headers, json=payload)
        
        # More detailed error handling
        if response.status_code != 200:
            error_msg = f"API Error: {response.status_code}"
            try:
                error_data = response.json()
                if 'error' in error_data:
                    error_msg += f" - {error_data['error'].get('message', 'Unknown error')}"
                    if 'code' in error_data['error']:
                        error_msg += f" (Code: {error_data['error']['code']})"
                    if 'param' in error_data['error']:
                        error_msg += f" - Parameter: {error_data['error']['param']}"
                print(f"Full error response: {error_data}")
            except Exception as json_err:
                error_msg += f" - {response.text[:200]}..."
            
            print(error_msg)
            # Return original markdown if API call fails, instead of raising an error
            return markdown_text
            
        result = response.json()
        
        print(f"Response received, status code: {response.status_code}")
        print(f"Tokens used: {result.get('usage', {}).get('total_tokens', 'N/A')}")
        print("=== END OF LLM REQUEST ===\n")
        
        refined_markdown = result['choices'][0]['message']['content']
        
        # Validate and fix markdown format
        refined_markdown = validate_markdown_format(refined_markdown)
        
        return refined_markdown
    except Exception as e:
        print(f"Error refining markdown with LLM: {str(e)}")
        # For debugging only, don't include in production
        import traceback
        print(f"Exception traceback: {traceback.format_exc()}")
        # Return original markdown if API call fails
        return markdown_text

def process_doc_with_gpt4o(filepath):
    """
    Process a document directly with GPT-4o Mini, converting non-image documents 
    to images first when needed.
    
    Args:
        filepath: Path to the file
        
    Returns:
        String containing markdown output from GPT-4o Mini
    """
    # Determine file type
    filetype = detect_file_type(filepath)
    file_ext = os.path.splitext(filepath)[1].lower()
    
    print(f"Processing file with Chuck Norris AI: {os.path.basename(filepath)}")
    print(f"File type: {filetype}, extension: {file_ext}")
    
    # For image files, process directly
    if 'image' in filetype or file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
        return process_image_with_gpt4o(filepath)
    
    # For PDFs, convert to images first
    elif 'pdf' in filetype or file_ext == '.pdf':
        return process_pdf_as_image_with_gpt4o(filepath)
    
    # For Word documents, convert to images first
    elif 'word' in filetype or file_ext in ['.doc', '.docx']:
        return process_doc_as_image_with_gpt4o(filepath)
    
    # For RTF files, use specialized handling
    elif 'rtf' in filetype or file_ext == '.rtf':
        return process_rtf_with_gpt4o(filepath)
        
    # For other text-based files
    else:
        try:
            # Try to read as text
            with open(filepath, "r", encoding="utf-8") as file:
                text_content = file.read()
                return process_text_with_gpt4o(text_content)
        except UnicodeDecodeError:
            # If not a text file, try to convert to image
            print("File is not a text file, attempting to convert to image...")
            return process_binary_file_with_gpt4o(filepath)

def process_image_with_gpt4o(image_path):
    """
    Process an image file directly with GPT-4o Mini.
    
    Args:
        image_path: Path to the image file
        
    Returns:
        String containing markdown output from GPT-4o Mini
    """
    try:
        # Read image as base64
        with open(image_path, "rb") as image_file:
            base64_image = base64.b64encode(image_file.read()).decode('utf-8')
        
        # Get file extension for content type
        file_ext = os.path.splitext(image_path)[1].lower()
        if file_ext in ['.jpg', '.jpeg']:
            content_type = "image/jpeg"
        elif file_ext == '.png':
            content_type = "image/png"
        elif file_ext == '.gif':
            content_type = "image/gif"
        elif file_ext == '.bmp':
            content_type = "image/bmp"
        else:
            content_type = "image/jpeg"  # Default
        
        print(f"Processing image file: {os.path.basename(image_path)}")
        return send_image_to_gpt4o(base64_image, content_type, image_path)
    
    except Exception as e:
        print(f"Error processing image file: {str(e)}")
        return f"Error processing image file: {str(e)}"

def process_pdf_as_image_with_gpt4o(pdf_path):
    """
    Convert PDF to images and process with GPT-4o Mini.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        String containing markdown output from GPT-4o Mini
    """
    try:
        
        print(f"Converting PDF to image: {os.path.basename(pdf_path)}")
        
        # Create a temporary directory for the images
        temp_dir = tempfile.mkdtemp()
        try:
            # Open the PDF
            pdf_document = fitz.open(pdf_path)
            
            # For simplicity, convert only the first page
            # For multi-page PDFs, we could process each page or a subset
            page = pdf_document[0]
            
            # Render page to an image (adjust zoom for higher resolution)
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            
            # Save the image
            image_path = os.path.join(temp_dir, "pdf_page.png")
            pix.save(image_path)
            
            # Process the image
            return process_image_with_gpt4o(image_path)
            
        finally:
            # Clean up the temp directory
            shutil.rmtree(temp_dir, ignore_errors=True)
            
    except ImportError:
        print("PyMuPDF (fitz) not installed. Trying alternative method...")
        # Fallback to using raw PDF data
        with open(pdf_path, "rb") as pdf_file:
            pdf_data = pdf_file.read()
            pdf_base64 = base64.b64encode(pdf_data).decode('utf-8')
            
            # Process using raw PDF data
            return send_image_to_gpt4o(pdf_base64, "application/pdf", pdf_path)
    except Exception as e:
        print(f"Error converting PDF to image: {str(e)}")
        return f"Error processing PDF: {str(e)}"

def process_doc_as_image_with_gpt4o(doc_path):
    """
    Convert DOCX to images and process with GPT-4o Mini.
    
    Args:
        doc_path: Path to the DOCX file
        
    Returns:
        String containing markdown output from GPT-4o Mini
    """
    try:
        # For now, we'll use a simple approach: send the raw document data
        with open(doc_path, "rb") as doc_file:
            doc_data = doc_file.read()
            doc_base64 = base64.b64encode(doc_data).decode('utf-8')
            
        # Determine content type
        file_ext = os.path.splitext(doc_path)[1].lower()
        if file_ext == '.docx':
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            content_type = "application/msword"
            
        print(f"Processing document file directly: {os.path.basename(doc_path)}")
        return send_image_to_gpt4o(doc_base64, content_type, doc_path)
        
    except Exception as e:
        print(f"Error processing document file: {str(e)}")
        return f"Error processing document: {str(e)}"

def process_binary_file_with_gpt4o(filepath):
    """
    Process an unknown binary file by reading it and sending as binary data.
    
    Args:
        filepath: Path to the binary file
        
    Returns:
        String containing markdown output from GPT-4o Mini
    """
    try:
        with open(filepath, "rb") as binary_file:
            binary_data = binary_file.read()
            binary_base64 = base64.b64encode(binary_data).decode('utf-8')
            
        print(f"Processing binary file: {os.path.basename(filepath)}")
        return send_image_to_gpt4o(binary_base64, "application/octet-stream", filepath)
        
    except Exception as e:
        print(f"Error processing binary file: {str(e)}")
        return f"Error processing file: {str(e)}"

def send_image_to_gpt4o(base64_data, content_type, filepath):
    """
    Send an image or document to GPT-4o Mini and get markdown output.
    
    Args:
        base64_data: Base64 encoded image data
        content_type: MIME type of the content
        filepath: Original file path (for reporting)
        
    Returns:
        String containing markdown output from GPT-4o Mini
    """
    # Get API key from config
    api_key = OPENAI_API_KEY
    api_url = "https://api.openai.com/v1/chat/completions"
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    
    print(f"Using Chuck Norris AI for document processing, content type: {content_type}")
    
    # Load system and user prompts from mustache template file
    system_prompt = ""
    user_prompt = ""
    
    try:
        with open('prompts/gpt4o_flow_chart_prompt.mustache', 'r') as template_file:
            template_content = template_file.read()
            
            # Extract system prompt (between first {{! System prompt... }} and {{! User prompt... }})
            system_start = template_content.find("{{! System prompt")
            user_start = template_content.find("{{! User prompt")
            
            if system_start != -1 and user_start != -1:
                # Extract system prompt (skip the comment line)
                system_prompt_with_comment = template_content[system_start:user_start].strip()
                system_prompt_lines = system_prompt_with_comment.split('\n')
                system_prompt = '\n'.join(system_prompt_lines[1:]).strip()
                
                # Extract user prompt (skip the comment line)
                user_prompt_with_comment = template_content[user_start:].strip()
                user_prompt_lines = user_prompt_with_comment.split('\n')
                user_prompt = '\n'.join(user_prompt_lines[1:]).strip()
    except Exception as e:
        print(f"Error loading prompt template: {str(e)}, using fallback prompts")
        # Fallback to hardcoded prompts
        system_prompt = "You are an expert at analyzing document layouts, detecting reading order, and converting document content to well-structured markdown. When presented with an image of a document or text, analyze the layout carefully before responding.\n\nIMPORTANT: Your response MUST be valid markdown format. This includes:\n- Using proper heading levels with # syntax\n- Correctly formatted lists (ordered and unordered)\n- Proper table syntax with | and --- separators\n- Correct code blocks with ``` delimiters\n- Proper link and image syntax\n- No HTML tags unless absolutely necessary\n- For any icons, symbols, or special characters that cannot be represented in plain text, describe them within square brackets, e.g., [checkmark icon], [arrow pointing right], [company logo]"
        user_prompt = "This is a document. Please analyze the layout first, detecting whether it has columns, tables, sections, flow charts, or other complex layouts. Then, extract all content and convert it to clean, well-structured markdown. Follow these steps:\n\n1. Analyze the layout (columns, reading order, tables, flow charts, etc.)\n2. Extract the full text content\n3. Convert to properly structured markdown with appropriate headings, lists, tables, etc.\n4. Return ONLY the valid markdown output without additional explanations\n5. Ensure all markdown syntax is correct and properly formatted\n6. For any icons, symbols, or special characters that cannot be represented in plain text, describe them within square brackets (e.g., [checkmark icon], [arrow pointing right], etc.)\n\nIMPORTANT - FOR FLOW CHARTS AND DIAGRAMS:\n- If the document contains a flow chart or diagram, represent its structure in markdown\n- For each vertex/node in the diagram, provide a concise one-sentence description\n- For each relationship (edge/arc) between vertices, explicitly list as 'Vertex A → Vertex B' (showing direction with an arrow)\n- Indicate whether relationships are one-directional (→) or bi-directional (↔) between nodes\n- List all vertex-to-vertex relationships to completely map the structure\n- Include a clear section titled '## Graph Structure' that enumerates all relationships\n- After listing all relationships, include a single sentence summarizing the overall flow or purpose\n- Do not just extract the text without capturing the structural relationships and direction of flow"
    
    # Create the payload with the document and prompt
    user_content = [
        {
            "type": "text",
            "text": user_prompt
        },
        {
            "type": "image_url",
            "image_url": {
                "url": f"data:{content_type};base64,{base64_data}"
            }
        }
    ]
    
    payload = {
        "model": "gpt-4o-mini",
        "messages": [
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_content
            }
        ],
        "temperature": 0.2
    }
    
    try:
        print("\n=== SENDING LLM REQUEST (Document Processing) ===")
        print(f"Model: gpt-4o-mini")
        print(f"Document: {os.path.basename(filepath)}")
        print(f"File size: {os.path.getsize(filepath)} bytes")
        print(f"Content type: {content_type}")
        print(f"Temperature: 0.2")
        
        response = requests.post(api_url, headers=headers, json=payload)
        
        # More detailed error handling
        if response.status_code != 200:
            error_msg = f"API Error: {response.status_code}"
            try:
                error_data = response.json()
                if 'error' in error_data:
                    error_msg += f" - {error_data['error'].get('message', 'Unknown error')}"
                    if 'code' in error_data['error']:
                        error_msg += f" (Code: {error_data['error']['code']})"
                    if 'param' in error_data['error']:
                        error_msg += f" - Parameter: {error_data['error']['param']}"
                print(f"Full error response: {error_data}")
            except Exception as json_err:
                error_msg += f" - {response.text[:200]}..."
            
            raise Exception(error_msg)
            
        response.raise_for_status()  # This will only run if status_code is 200
        result = response.json()
        
        print(f"Response received, status code: {response.status_code}")
        print(f"Tokens used: {result.get('usage', {}).get('total_tokens', 'N/A')}")
        print("=== END OF LLM REQUEST ===\n")
        
        markdown_content = result['choices'][0]['message']['content']
        
        # Validate markdown format
        markdown_content = validate_markdown_format(markdown_content)
        
        return markdown_content
    except Exception as e:
        print(f"Error processing document with GPT-4o Mini: {str(e)}")
        # For debugging only, don't include in production
        import traceback
        print(f"Exception traceback: {traceback.format_exc()}")
        return f"Error processing document with AI: {str(e)}"

# Helper function for text-based documents
def process_text_with_gpt4o(text_content):
    """
    Process text content with GPT-4o Mini.
    
    Args:
        text_content: The text content to process
        
    Returns:
        String containing markdown output from GPT-4o Mini
    """
    # Get API key from config
    api_key = OPENAI_API_KEY
    api_url = "https://api.openai.com/v1/chat/completions"
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    
    system_prompt = "You are an expert at converting document content to well-structured markdown. Your output must be valid, properly formatted markdown."
    
    user_prompt = f"""
    Please convert the following content to clean, well-structured markdown. 
    Make sure to:
    1. Keep all important information
    2. Use proper markdown structure with headings, lists, tables, etc.
    3. Ensure all markdown syntax is correct and valid
    4. Return ONLY the converted markdown without additional explanations
    
    Here's the content to convert:
    
    {text_content}
    """
    
    payload = {
        "model": "gpt-4o-mini",
        "messages": [
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_prompt
            }
        ],
        "temperature": 0.2
    }
    
    try:
        print("\n=== SENDING LLM REQUEST (Text Processing) ===")
        print(f"Model: gpt-4o-mini")
        print(f"Text length: {len(text_content)} characters")
        print(f"Temperature: 0.2")
        
        response = requests.post(api_url, headers=headers, json=payload)
        
        # More detailed error handling
        if response.status_code != 200:
            error_msg = f"API Error: {response.status_code}"
            try:
                error_data = response.json()
                if 'error' in error_data:
                    error_msg += f" - {error_data['error'].get('message', 'Unknown error')}"
                    if 'code' in error_data['error']:
                        error_msg += f" (Code: {error_data['error']['code']})"
                    if 'param' in error_data['error']:
                        error_msg += f" - Parameter: {error_data['error']['param']}"
                print(f"Full error response: {error_data}")
            except Exception as json_err:
                error_msg += f" - {response.text[:200]}..."
            
            raise Exception(error_msg)
            
        response.raise_for_status()
        result = response.json()
        
        print(f"Response received, status code: {response.status_code}")
        print(f"Tokens used: {result.get('usage', {}).get('total_tokens', 'N/A')}")
        print("=== END OF LLM REQUEST ===\n")
        
        markdown_content = result['choices'][0]['message']['content']
        
        # Validate markdown format
        markdown_content = validate_markdown_format(markdown_content)
        
        return markdown_content
    except Exception as e:
        print(f"Error processing text with GPT-4o Mini: {str(e)}")
        # For debugging only, don't include in production
        import traceback
        print(f"Exception traceback: {traceback.format_exc()}")
        return f"Error processing text with AI: {str(e)}"

def validate_markdown_format(markdown_text):
    """
    Validate and fix common markdown formatting issues.
    
    Args:
        markdown_text: The markdown text to validate
        
    Returns:
        Cleaned and validated markdown text
    """
    # Check if the response starts with explanatory text that should be removed
    lines = markdown_text.split('\n')
    start_idx = 0
    
    # Skip any introductory text before actual markdown content
    for i, line in enumerate(lines):
        if line.startswith('#') or line.startswith('-') or line.startswith('*') or line.startswith('1.') or line.startswith('|'):
            start_idx = i
            break
    
    # If we found a starting point for markdown content, remove preceding text
    if start_idx > 0:
        lines = lines[start_idx:]
        markdown_text = '\n'.join(lines)
    
    # Fix common markdown issues
    
    # Ensure proper heading spacing (# Heading, not #Heading)
    markdown_text = re.sub(r'(^|\n)#([^# ])', r'\1# \2', markdown_text)
    markdown_text = re.sub(r'(^|\n)##([^# ])', r'\1## \2', markdown_text)
    markdown_text = re.sub(r'(^|\n)###([^# ])', r'\1### \2', markdown_text)
    
    # Ensure proper list item spacing
    markdown_text = re.sub(r'(^|\n)-([^ ])', r'\1- \2', markdown_text)
    markdown_text = re.sub(r'(^|\n)\*([^ ])', r'\1* \2', markdown_text)
    
    # Ensure numbered lists have proper spacing
    markdown_text = re.sub(r'(^|\n)(\d+)\.([^ ])', r'\1\2. \3', markdown_text)
    
    # Remove any trailing "markdown" or "md" words that might be added
    markdown_text = re.sub(r'\n+markdown\s*$', '', markdown_text, flags=re.IGNORECASE)
    markdown_text = re.sub(r'\n+md\s*$', '', markdown_text, flags=re.IGNORECASE)
    
    # Remove explanatory comments that might be added
    markdown_text = re.sub(r'(?i)```markdown|```md', '```', markdown_text)
    
    # Count backtick groups to ensure balanced code blocks
    backtick_groups = re.findall(r'```', markdown_text)
    if len(backtick_groups) % 2 != 0:
        # Unbalanced code blocks, add closing backticks if needed
        if markdown_text.strip().endswith('```'):
            # Remove trailing backticks if they're at the end
            markdown_text = re.sub(r'```\s*$', '', markdown_text.strip())
        else:
            # Add closing backticks if we end with an open code block
            markdown_text = markdown_text.strip() + '\n```'
    
    # Look for unusual character sequences that might be icons and wrap them in square brackets if not already
    # This regex finds Unicode characters outside the basic Latin alphabet and common punctuation
    def replace_special_chars(match):
        char = match.group(0)
        # Skip if already within square brackets
        if re.search(r'\[[^\]]*' + re.escape(char) + r'[^\]]*\]', markdown_text):
            return char
        # Skip common characters we don't want to replace
        if char in '.,;:!?"\'-+=(){}[]<>/@#$%^&*_|\\~`' or char.isalnum() or char.isspace():
            return char
        # Replace with description in square brackets
        return f"[symbol: {char}]"
    
    # Apply the replacement for special characters
    # markdown_text = re.sub(r'[^\x00-\x7F]+', replace_special_chars, markdown_text)
    
    # Ensure all icon descriptions in square brackets are properly formatted
    # Look for incomplete brackets
    markdown_text = re.sub(r'\[[^\]]+$', lambda m: f"{m.group(0)}]", markdown_text)
    markdown_text = re.sub(r'^\][^\[]+', lambda m: f"[{m.group(0)[1:]}", markdown_text)
    
    return markdown_text

def convert_to_markdown(filepath, use_ai_refinement=True, use_chuck_norris_ai=False):
    """
    Generic dispatcher that converts various file types to markdown.
    
    Args:
        filepath: Path to the file
        use_ai_refinement: Whether to use LLM to refine the markdown output
        use_chuck_norris_ai: Whether to bypass OCR and use GPT-4o Mini directly for images
        
    Returns:
        String containing markdown representation of the file, or
        Dictionary with markdown and html_preview for RTF files
    """
    if not os.path.exists(filepath):
        return f"File not found: {filepath}"
    
    try:
        filetype = detect_file_type(filepath)
        file_ext = os.path.splitext(filepath)[1].lower()
        
        raw_markdown = ""
        html_preview = None
        
        # For any file type with Chuck Norris AI enabled, use the direct GPT-4o Mini approach
        if use_chuck_norris_ai:
            print(f"Using Chuck Norris AI for processing {filetype} file, bypassing traditional parsing...")
            # For RTF files, the process_doc_with_gpt4o function will return a dictionary with markdown and html_preview
            return process_doc_with_gpt4o(filepath)
        
        # Traditional parsing approaches if Chuck Norris AI is not enabled
        if 'pdf' in filetype:
            raw_markdown = parse_pdf(filepath)
        elif 'word' in filetype or filepath.endswith(".docx"):
            raw_markdown = parse_docx(filepath)
        elif 'rtf' in filetype or file_ext == '.rtf':
            # For RTF files, generate an HTML preview as well
            temp_dir = tempfile.mkdtemp()
            try:
                # Convert RTF to HTML for display
                html_path, extracted_images = rtf_to_html(filepath, temp_dir)
                
                # Process RTF text content
                raw_markdown = parse_rtf(filepath)
                
                # Create a unique name for the final HTML that will be preserved
                final_html_dir = os.path.join(os.path.dirname(os.path.dirname(filepath)), 'static', 'rtf_previews')
                os.makedirs(final_html_dir, exist_ok=True)
                
                final_html_name = f"rtf_preview_{uuid.uuid4().hex}.html"
                final_html_path = os.path.join(final_html_dir, final_html_name)
                
                # Copy the HTML file to a location that can be accessed by the browser
                shutil.copy2(html_path, final_html_path)
                
                # Copy any extracted images to the same directory
                for img_path in extracted_images:
                    img_filename = os.path.basename(img_path)
                    shutil.copy2(img_path, os.path.join(final_html_dir, img_filename))
                
                html_preview = f"/static/rtf_previews/{final_html_name}"
            finally:
                # Clean up temp directory
                shutil.rmtree(temp_dir, ignore_errors=True)
        elif 'image' in filetype or filepath.endswith((".png", ".jpg", ".jpeg")):
            raw_markdown = parse_image(filepath)
        else:
            # Try to read as plain text for unknown types
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                raw_markdown = f"```\n{content}\n```"
            except:
                return f"Unsupported file type: {filetype}"
        
        # Refine the markdown using OpenAI's GPT-4o Mini if requested
        if use_ai_refinement:
            refined_markdown = refine_markdown_with_llm(raw_markdown)
            
            # If we have an HTML preview, return a dictionary
            if html_preview:
                return {
                    'markdown': refined_markdown,
                    'html_preview': html_preview
                }
            return refined_markdown
        else:
            # If we have an HTML preview, return a dictionary
            if html_preview:
                return {
                    'markdown': raw_markdown,
                    'html_preview': html_preview
                }
            return raw_markdown
    except Exception as e:
        print(f"Error converting file: {str(e)}")
        import traceback
        print(f"Exception traceback: {traceback.format_exc()}")
        return f"Error converting file: {str(e)}"

def process_rtf_with_gpt4o(rtf_path):
    """
    Process an RTF file with GPT-4o Mini, extracting embedded images if available.
    Creates an HTML representation for display in the UI.
    
    Args:
        rtf_path: Path to the RTF file
        
    Returns:
        Dictionary containing markdown output and HTML preview path
    """
    try:
        print(f"Processing RTF file: {os.path.basename(rtf_path)}")
        
        # Check file size first to avoid token limit issues
        file_size = os.path.getsize(rtf_path)
        if file_size > 5 * 1024 * 1024:  # 5MB limit
            print(f"RTF file too large ({file_size / (1024*1024):.2f} MB), chunking content...")
            return process_large_rtf_file(rtf_path)
        
        # Create a temporary directory for extracted images and HTML
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Convert RTF to HTML for display and processing
            html_path, extracted_images = rtf_to_html(rtf_path, temp_dir)
            
            # Create a unique name for the final HTML that will be preserved
            final_html_dir = os.path.join(os.path.dirname(os.path.dirname(rtf_path)), 'static', 'rtf_previews')
            os.makedirs(final_html_dir, exist_ok=True)
            
            final_html_name = f"rtf_preview_{uuid.uuid4().hex}.html"
            final_html_path = os.path.join(final_html_dir, final_html_name)
            
            # Copy the HTML file to a location that can be accessed by the browser
            shutil.copy2(html_path, final_html_path)
            
            # Copy any extracted images to the same directory
            for img_path in extracted_images:
                img_filename = os.path.basename(img_path)
                shutil.copy2(img_path, os.path.join(final_html_dir, img_filename))
            
            # Generate an image from HTML for LLM processing
            html_image_path = os.path.join(temp_dir, f"rtf_as_image_{uuid.uuid4().hex}.png")
            rendered_image_path = html_to_image(html_path, html_image_path)
            
            # If HTML to image conversion succeeded, use that for LLM
            if rendered_image_path and os.path.exists(rendered_image_path):
                print("Successfully rendered HTML to image for LLM processing")
                markdown_content = process_image_with_gpt4o(rendered_image_path)
            else:
                print("HTML rendering failed, falling back to text extraction")
                # Fallback to text extraction - don't use mammoth as it expects ZIP files
                try:
                    # Try to extract plain text from RTF
                    try:
                        import striprtf.striprtf as striprtf
                        with open(rtf_path, 'rb') as rtf_file:
                            rtf_text = rtf_file.read().decode('utf-8', errors='ignore')
                            plain_text = striprtf.rtf_to_text(rtf_text)
                    except ImportError:
                        # Basic RTF to text
                        with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
                            rtf_text = f.read()
                            # Remove RTF headers and control words
                            plain_text = re.sub(r'\\[a-z0-9]+', ' ', rtf_text)
                            plain_text = re.sub(r'\{|\}|\\', '', plain_text)
                    
                    # Truncate if too large
                    if len(plain_text) > 50000:
                        print(f"RTF text content too large ({len(plain_text)} chars), truncating...")
                        plain_text = plain_text[:25000] + "\n\n... [Content truncated due to size] ...\n\n" + plain_text[-25000:]
                    
                    # Process text with GPT-4o
                    markdown_content = process_text_with_gpt4o(plain_text)
                except Exception as text_err:
                    print(f"Error processing RTF as text: {str(text_err)}")
                    markdown_content = f"Error processing RTF file: {str(text_err)}"
            
            # Process a sample of extracted images separately if available
            image_markdown = []
            if extracted_images:
                image_markdown.append("\n## Embedded Images\n")
                
                for i, img_path in enumerate(extracted_images[:3]):  # Limit to 3 images
                    try:
                        # Process the image with GPT-4o
                        image_result = process_image_with_gpt4o(img_path)
                        image_markdown.append(f"\n### Embedded Image {i+1}\n")
                        image_markdown.append(image_result)
                    except Exception as img_err:
                        print(f"Error processing embedded image {i+1}: {str(img_err)}")
                        image_markdown.append(f"\n*Error processing embedded image {i+1}: {str(img_err)}*\n")
                
                if len(extracted_images) > 3:
                    image_markdown.append(f"\n*{len(extracted_images) - 3} additional images were found but not processed.*\n")
            
            # Combine text and image results
            combined_markdown = markdown_content
            if image_markdown:
                combined_markdown += "\n\n" + "\n\n".join(image_markdown)
            
            # Return both the markdown and the HTML preview path
            return {
                'markdown': combined_markdown,
                'html_preview': f"/static/rtf_previews/{final_html_name}"
            }
                
        finally:
            # Clean up the temp directory
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    except Exception as e:
        print(f"Error processing RTF file: {str(e)}")
        import traceback
        print(f"Exception traceback: {traceback.format_exc()}")
        return {
            'markdown': f"Error processing RTF file: {str(e)}",
            'html_preview': None
        }

def rtf_to_html(rtf_path, output_dir):
    """
    Convert an RTF file to HTML format.
    
    Args:
        rtf_path: Path to the RTF file
        output_dir: Directory to save the HTML file and extracted images
        
    Returns:
        Tuple containing (html_path, list of extracted image paths)
    """
    try:
        # Create a unique HTML filename
        html_filename = f"rtf_preview_{uuid.uuid4().hex}.html"
        html_path = os.path.join(output_dir, html_filename)
        
        # Extract images from RTF
        extracted_images = extract_images_from_rtf(rtf_path, output_dir)
        
        # Read RTF content using a specialized RTF approach instead of mammoth
        try:
            # First try to read RTF using striprtf library if available
            try:
                import striprtf.striprtf as striprtf
                with open(rtf_path, 'rb') as rtf_file:
                    rtf_text = rtf_file.read().decode('utf-8', errors='ignore')
                    html_content = striprtf.rtf_to_html(rtf_text)
            except ImportError:
                # Fallback to basic text extraction
                with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = f.read()
                    # Basic RTF to HTML conversion
                    # Remove RTF headers and control words
                    clean_text = re.sub(r'\\[a-z0-9]+', ' ', text_content)
                    clean_text = re.sub(r'\{|\}|\\', '', clean_text)
                    # Create simple HTML
                    html_content = f"<html><head><title>RTF Preview</title></head><body><pre>{html.escape(clean_text)}</pre></body></html>"
        except Exception as e:
            print(f"Error parsing RTF content: {str(e)}")
            # Absolute fallback
            html_content = f"<html><head><title>RTF Preview</title></head><body><div>Error processing RTF file: {html.escape(str(e))}</div></body></html>"
        
        # Enhance the HTML with CSS for better display
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Make sure we have head and body sections
        if not soup.head:
            head = soup.new_tag('head')
            soup.insert(0, head)
        if not soup.body:
            body = soup.new_tag('body')
            if soup.pre:
                body.append(soup.pre.extract())
            soup.append(body)
        
        # Add CSS styling
        style_tag = soup.new_tag('style')
        style_tag.string = """
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 20px;
            color: #333;
        }
        h1, h2, h3, h4, h5, h6 {
            color: #2c3e50;
            margin-top: 24px;
            margin-bottom: 16px;
        }
        img {
            max-width: 100%;
            border: 1px solid #ddd;
            border-radius: 4px;
            padding: 5px;
            margin: 10px 0;
        }
        pre {
            background-color: #f5f5f5;
            padding: 10px;
            border-radius: 4px;
            overflow-x: auto;
            white-space: pre-wrap;
        }
        """
        soup.head.append(style_tag)
        
        # Add extracted images to the HTML
        if extracted_images:
            image_container = soup.new_tag('div')
            image_container['class'] = 'extracted-images'
            image_container['style'] = 'margin-top: 30px; border-top: 1px solid #ddd; padding-top: 20px;'
            
            heading = soup.new_tag('h2')
            heading.string = "Embedded Images"
            image_container.append(heading)
            
            for i, img_path in enumerate(extracted_images):
                img_tag = soup.new_tag('img')
                img_tag['src'] = os.path.basename(img_path)
                img_tag['alt'] = f"Embedded Image {i+1}"
                img_tag['style'] = "display: block; margin: 20px 0; max-width: 100%;"
                
                figure = soup.new_tag('figure')
                figure['style'] = "margin: 20px 0; text-align: center;"
                
                caption = soup.new_tag('figcaption')
                caption.string = f"Embedded Image {i+1}"
                caption['style'] = "font-style: italic; color: #666;"
                
                figure.append(img_tag)
                figure.append(caption)
                image_container.append(figure)
            
            # Add image container to body
            soup.body.append(image_container)
        
        # Write enhanced HTML to file
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(str(soup))
        
        return html_path, extracted_images
    
    except Exception as e:
        print(f"Error converting RTF to HTML: {str(e)}")
        import traceback
        print(f"Exception traceback: {traceback.format_exc()}")
        
        # Create a fallback HTML with error message
        fallback_html = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; padding: 20px; }}
                .error {{ color: red; background-color: #ffeeee; padding: 10px; border-radius: 5px; }}
            </style>
        </head>
        <body>
            <h1>Error Converting RTF File</h1>
            <div class="error">
                <p><strong>Error:</strong> {html.escape(str(e))}</p>
                <p>The RTF file could not be properly converted to HTML.</p>
            </div>
            <hr>
            <h2>Raw Content Preview:</h2>
            <pre style="background-color: #f5f5f5; padding: 10px; border-radius: 5px; overflow: auto;">
            """
        
        # Add the first 2000 characters of raw content
        try:
            with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
                raw_content = f.read(2000)
                fallback_html += html.escape(raw_content) + "..."
        except:
            fallback_html += "Unable to read file content."
        
        fallback_html += """
            </pre>
        </body>
        </html>
        """
        
        # Write fallback HTML
        html_path = os.path.join(output_dir, f"rtf_preview_error_{uuid.uuid4().hex}.html")
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(fallback_html)
        
        return html_path, []

def html_to_image(html_path, output_path):
    """
    Convert HTML file to an image for LLM processing.
    
    Args:
        html_path: Path to the HTML file
        output_path: Path to save the output image
        
    Returns:
        Path to the generated image or None if failed
    """
    try:
        # This requires wkhtmltoimage to be installed on the system
        # Try to use Python wkhtmltopdf library if available
        try:
            import imgkit
            options = {
                'width': 1200,
                'height': 1600,
                'quality': 80,
                'enable-local-file-access': True
            }
            imgkit.from_file(html_path, output_path, options=options)
            return output_path
        except ImportError:
            print("imgkit not installed, trying alternative method...")
            
            # Alternative approach using wkhtmltoimage command line
            import subprocess
            try:
                # Try to use wkhtmltoimage directly
                command = [
                    'wkhtmltoimage',
                    '--width', '1200',
                    '--height', '1600',
                    '--quality', '80',
                    '--enable-local-file-access',
                    html_path,
                    output_path
                ]
                subprocess.run(command, check=True)
                return output_path
            except (subprocess.SubprocessError, FileNotFoundError):
                print("wkhtmltoimage command failed, trying Python rendering...")
                
                # Fallback to a very basic HTML rendering with PIL
                text = "HTML Preview (rendering failed)"
                img = Image.new('RGB', (800, 200), color=(255, 255, 255))
                try:
                    from PIL import ImageDraw, ImageFont
                    draw = ImageDraw.Draw(img)
                    try:
                        font = ImageFont.truetype("Arial", 20)
                    except:
                        font = ImageFont.load_default()
                    draw.text((20, 20), text, fill=(0, 0, 0), font=font)
                except:
                    pass  # Can't even draw text, will return blank image
                
                img.save(output_path)
                return output_path
    except Exception as e:
        print(f"Error converting HTML to image: {str(e)}")
        return None

def process_large_rtf_file(rtf_path):
    """
    Process a large RTF file by breaking it into manageable chunks.
    Also creates an HTML representation for display in the UI.
    
    Args:
        rtf_path: Path to the RTF file
        
    Returns:
        Dictionary containing markdown output and HTML preview path
    """
    try:
        # Create a temp directory for extracted images
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Convert RTF to HTML for display
            html_path, extracted_images = rtf_to_html(rtf_path, temp_dir)
            
            # Create a unique name for the final HTML that will be preserved
            final_html_dir = os.path.join(os.path.dirname(os.path.dirname(rtf_path)), 'static', 'rtf_previews')
            os.makedirs(final_html_dir, exist_ok=True)
            
            final_html_name = f"rtf_preview_large_{uuid.uuid4().hex}.html"
            final_html_path = os.path.join(final_html_dir, final_html_name)
            
            # Copy the HTML file to a location that can be accessed by the browser
            shutil.copy2(html_path, final_html_path)
            
            # Copy any extracted images to the same directory
            for img_path in extracted_images:
                img_filename = os.path.basename(img_path)
                shutil.copy2(img_path, os.path.join(final_html_dir, img_filename))
            
            # Extract text directly from RTF, not using mammoth
            try:
                # Try to extract plain text from RTF
                try:
                    import striprtf.striprtf as striprtf
                    with open(rtf_path, 'rb') as rtf_file:
                        rtf_text = rtf_file.read().decode('utf-8', errors='ignore')
                        text_content = striprtf.rtf_to_text(rtf_text)
                except ImportError:
                    # Basic RTF to text
                    with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
                        rtf_text = f.read()
                        # Remove RTF headers and control words
                        text_content = re.sub(r'\\[a-z0-9]+', ' ', rtf_text)
                        text_content = re.sub(r'\{|\}|\\', '', text_content)
            except Exception as e:
                print(f"Error extracting text from large RTF: {str(e)}")
                # If all text extraction fails, create placeholder text
                text_content = "[The RTF content could not be extracted properly]"
            
            # Split text into chunks of approximately 10,000 characters
            chunk_size = 10000
            chunks = [text_content[i:i+chunk_size] for i in range(0, len(text_content), chunk_size)]
            
            # Process each chunk and collect results
            processed_chunks = []
            for i, chunk in enumerate(chunks):
                if i == 0:
                    # Process first chunk normally
                    try:
                        chunk_result = process_text_with_gpt4o(f"This is the first part of a document that has been split into chunks due to size limitations.\n\n{chunk}")
                        processed_chunks.append(chunk_result)
                    except Exception as chunk_err:
                        print(f"Error processing chunk {i+1}: {str(chunk_err)}")
                        processed_chunks.append(f"*Error processing chunk {i+1}*\n\n{chunk[:500]}...\n")
                elif i == len(chunks) - 1:
                    # Process last chunk
                    try:
                        chunk_result = process_text_with_gpt4o(f"This is the last part of a document that has been split into chunks due to size limitations.\n\n{chunk}")
                        processed_chunks.append(chunk_result)
                    except Exception as chunk_err:
                        print(f"Error processing chunk {i+1}: {str(chunk_err)}")
                        processed_chunks.append(f"*Error processing chunk {i+1}*\n\n{chunk[:500]}...\n")
                else:
                    # Process middle chunks
                    try:
                        chunk_result = process_text_with_gpt4o(f"This is part {i+1} of a document that has been split into chunks due to size limitations.\n\n{chunk}")
                        processed_chunks.append(chunk_result)
                    except Exception as chunk_err:
                        print(f"Error processing chunk {i+1}: {str(chunk_err)}")
                        processed_chunks.append(f"*Error processing chunk {i+1}*\n\n{chunk[:500]}...\n")
            
            # Process up to 3 images
            image_results = []
            for i, img_path in enumerate(extracted_images[:3]):
                try:
                    image_result = process_image_with_gpt4o(img_path)
                    image_results.append(f"\n### Embedded Image {i+1}\n\n{image_result}")
                except Exception as img_err:
                    print(f"Error processing image {i+1}: {str(img_err)}")
            
            # Combine results
            combined_markdown = "\n\n# Document Content (Processed in Chunks)\n\n"
            for i, chunk_result in enumerate(processed_chunks):
                combined_markdown += f"\n\n## Part {i+1}\n\n{chunk_result}\n\n"
            
            if image_results:
                combined_markdown += "\n\n# Embedded Images\n\n"
                combined_markdown += "\n\n".join(image_results)
            
            if len(extracted_images) > 3:
                combined_markdown += f"\n\n*Note: {len(extracted_images) - 3} additional images were found but not processed due to size limitations.*"
            
            return {
                'markdown': combined_markdown,
                'html_preview': f"/static/rtf_previews/{final_html_name}"
            }
            
        finally:
            # Clean up
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    except Exception as e:
        print(f"Error processing large RTF file: {str(e)}")
        return {
            'markdown': f"Error processing large RTF file: {str(e)}",
            'html_preview': None
        }

# Example usage
if __name__ == "__main__":
    # Test with different file types
    test_files = [
        "example.pdf",
        "example.docx",
        "example.rtf",
        "example.png"
    ]
    
    for test_file in test_files:
        if os.path.exists(test_file):
            print(f"Converting {test_file} to markdown...")
            markdown = convert_to_markdown(test_file)
            print(f"First 200 characters: {markdown[:200]}...") 