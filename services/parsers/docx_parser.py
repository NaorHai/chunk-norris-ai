import os
import tempfile
import shutil
from typing import List, Optional, Dict, Any

from docx import Document

from services.parsers.file_parser_interface import FileParserInterface


class DOCXParser(FileParserInterface):
    """
    Parser for DOCX files that converts them to markdown
    """
    
    def __init__(self, output_dir: Optional[str] = None):
        """
        Initialize the DOCX parser
        
        Args:
            output_dir: Optional directory to save extracted images
        """
        self.output_dir = output_dir
        
    def parse(self, filepath: str) -> str:
        """
        Parse a DOCX file and convert it to markdown format.
        
        Args:
            filepath: Path to the DOCX file
            
        Returns:
            String containing markdown representation of the DOCX
        """
        try:
            doc = Document(filepath)
            md = []
            
            # Create a temp directory for extracted images
            temp_dir = tempfile.mkdtemp()
            try:
                # Extract images if present (this requires additional handling)
                image_count = 0
                
                for p in doc.paragraphs:
                    if p.style.name.startswith("Heading"):
                        level = p.style.name[-1]
                        md.append(f"{'#' * int(level)} {p.text}")
                    else:
                        md.append(p.text)
                
                # Optional: Process tables from DOCX
                for table in doc.tables:
                    md.append("\n### Table\n")
                    table_rows = []
                    
                    # Process header row
                    header_row = []
                    for cell in table.rows[0].cells:
                        header_row.append(cell.text.strip())
                    
                    table_rows.append("| " + " | ".join(header_row) + " |")
                    table_rows.append("| " + " | ".join(["---"] * len(header_row)) + " |")
                    
                    # Process data rows
                    for row in table.rows[1:]:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        table_rows.append("| " + " | ".join(row_data) + " |")
                    
                    md.append("\n".join(table_rows))
            
            finally:
                # Clean up the temp directory
                shutil.rmtree(temp_dir, ignore_errors=True)
            
            return "\n\n".join(md)
        
        except Exception as e:
            return f"Error parsing DOCX: {str(e)}" 